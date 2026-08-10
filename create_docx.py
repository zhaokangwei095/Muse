from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE

OUT = r"D:\桌面\muse\muse_灵感探索功能说明.docx"

def shade(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr(); shd = OxmlElement('w:shd'); shd.set(qn('w:fill'), fill); tcPr.append(shd)

def set_cell_border(cell, color='D8DDE8', sz='8'):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr(); borders = tcPr.first_child_found_in('w:tcBorders')
    if borders is None: borders = OxmlElement('w:tcBorders'); tcPr.append(borders)
    for edge in ('top','left','bottom','right','insideH','insideV'):
        tag='w:'+edge; el=borders.find(qn(tag))
        if el is None: el=OxmlElement(tag); borders.append(el)
        el.set(qn('w:val'),'single'); el.set(qn('w:sz'),sz); el.set(qn('w:color'),color)

def add_placeholder(doc, label='图片占位｜请插入产品截图或流程图', height=2.2):
    t=doc.add_table(rows=1, cols=1); t.alignment=WD_TABLE_ALIGNMENT.CENTER; t.autofit=False
    c=t.cell(0,0); c.width=Inches(6.35); c.height=Inches(height); c.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER; shade(c,'F3F5F8'); set_cell_border(c)
    p=c.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=p.add_run(label); r.font.size=Pt(11); r.font.color.rgb=RGBColor(120,130,148)
    doc.add_paragraph()

def add_bullets(doc, items):
    for x in items:
        p=doc.add_paragraph(style='List Bullet'); p.paragraph_format.space_after=Pt(4); p.add_run(x)

def add_section(doc, no, title, intro, image_label, points, flow):
    doc.add_heading(f'{no:02d}  {title}', level=1)
    p=doc.add_paragraph(intro); p.style='Body Text'
    add_placeholder(doc, image_label)
    doc.add_heading('功能要点', level=2); add_bullets(doc, points)
    doc.add_heading('用户路径', level=2); doc.add_paragraph(flow)

doc=Document(); sec=doc.sections[0]; sec.top_margin=Inches(.7); sec.bottom_margin=Inches(.7); sec.left_margin=Inches(.85); sec.right_margin=Inches(.85)
styles=doc.styles
styles['Normal'].font.name='Microsoft YaHei'; styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'),'Microsoft YaHei'); styles['Normal'].font.size=Pt(10.5); styles['Normal'].font.color.rgb=RGBColor(54,63,78)
for s in ('Title','Heading 1','Heading 2'):
    styles[s].font.name='Microsoft YaHei'; styles[s]._element.rPr.rFonts.set(qn('w:eastAsia'),'Microsoft YaHei')
styles['Title'].font.size=Pt(30); styles['Title'].font.bold=True; styles['Title'].font.color.rgb=RGBColor(39,51,76)
styles['Heading 1'].font.size=Pt(19); styles['Heading 1'].font.bold=True; styles['Heading 1'].font.color.rgb=RGBColor(39,51,76)
styles['Heading 2'].font.size=Pt(12.5); styles['Heading 2'].font.bold=True; styles['Heading 2'].font.color.rgb=RGBColor(91,105,133)

p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.space_after=Pt(20)
r=p.add_run('MUSE'); r.font.size=Pt(15); r.font.bold=True; r.font.color.rgb=RGBColor(112,125,153)
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=p.add_run('沉浸式灵感探索'); r.font.size=Pt(30); r.font.bold=True; r.font.color.rgb=RGBColor(39,51,76)
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=p.add_run('功能介绍与体验说明'); r.font.size=Pt(15); r.font.color.rgb=RGBColor(112,125,153)
doc.add_paragraph('\n')
add_placeholder(doc,'封面视觉占位｜插入品牌主视觉 / 产品界面拼图',3.0)
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=p.add_run('版本 1.0  ·  2026'); r.font.size=Pt(9); r.font.color.rgb=RGBColor(145,153,168)
doc.add_page_break()

doc.add_heading('体验概览', level=1)
doc.add_paragraph('MUSE 将灵感发现、情绪探索与创作者交流整合在同一条体验链路中。用户可以轻松浏览作品、表达偏好、收藏沉淀，并在合适的时刻与创作者展开对话。整体体验强调轻盈、沉浸和可持续探索。')
tbl=doc.add_table(rows=1, cols=3); tbl.alignment=WD_TABLE_ALIGNMENT.CENTER
for i,(a,b) in enumerate([('发现','卡片滑动 / 瀑布流'),('连接','创作者聊天 / AI辅助'),('沉淀','灵感收藏 / 热度数据')]):
    c=tbl.cell(0,i); shade(c,'EEF1F7'); set_cell_border(c); p=c.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER; rr=p.add_run(a+'\n'); rr.bold=True; rr.font.color.rgb=RGBColor(39,51,76); p.add_run(b).font.size=Pt(9)
doc.add_paragraph(); doc.add_heading('设计关键词', level=2); doc.add_paragraph('半透明 · 柔和光影 · 圆润卡片 · 明暗主题 · 情绪化反馈 · 低门槛互动')
doc.add_page_break()

add_section(doc,1,'沉浸式灵感探索','通过卡片滑动浏览灵感内容，左右滑动即可表达喜欢或跳过，也可以进入详情页深入了解。移动端支持卡片模式与瀑布流模式切换，让探索节奏由用户掌控。','图片占位｜卡片模式、瀑布流模式、详情页', ['左右滑动：喜欢 / 跳过，形成连续探索节奏','点击卡片：进入作品详情，查看内容与创作者信息','模式切换：卡片模式适合沉浸浏览，瀑布流适合快速扫选','操作反馈：收藏、跳过与切换均提供即时动效或状态变化'], '打开探索页 → 选择浏览模式 → 浏览作品 → 喜欢 / 跳过 → 进入详情或继续探索')
add_section(doc,2,'沉浸式视觉设计','以半透明材质、柔和光影与圆润卡片建立统一的高级感。通过明暗主题切换，适配不同时间与环境下的浏览状态。','图片占位｜明亮主题与暗色主题对比', ['半透明层次：区分背景、内容与操作层级','柔和光影：降低视觉噪声，强化内容聚焦','圆润卡片：传递友好、轻盈与可触感','明暗主题：一键切换，并保持信息层级和对比度一致'], '进入设置 / 顶部主题入口 → 切换明亮或暗色 → 页面即时更新 → 下次访问保留偏好')
add_section(doc,3,'创作者聊天','浏览作品时可进入与创作者的聊天空间，围绕作品想法和创作灵感展开交流，让发现从单向浏览延伸为双向互动。','图片占位｜聊天空间、回复状态、AI辅助回复', ['消息发送：支持文字与图片等常用内容形式','回复状态：展示发送中、已送达与对方正在输入等状态','AI辅助回复：根据当前作品与上下文生成回复建议','上下文衔接：从作品详情进入聊天时，保留作品引用或预览'], '作品详情 → 点击联系创作者 → 输入消息 / 选择AI建议 → 发送 → 查看回复状态')
add_section(doc,4,'灵感收藏','用户可以一键收藏喜欢的作品，并在个人空间集中管理自己的灵感库。通过内容类型筛选，快速找到适合当下需求的参考。','图片占位｜收藏按钮、个人灵感库、类型筛选', ['一键收藏：在卡片和详情页均可完成','集中查看：个人空间展示收藏内容与最近更新','类型筛选：按插画、摄影、设计、文字等类型浏览','取消收藏：支持在内容卡片或详情页快速移除'], '看到喜欢的作品 → 点击收藏 → 进入个人空间 → 选择类型筛选 → 查看或取消收藏')
add_section(doc,5,'塔罗式情绪探索','将传统塔罗的“抽取与解读”转化为现代情绪探索。用户经过简短呼吸引导后抽取一张专属情绪卡，获得对当前状态的理解与下一步灵感。','图片占位｜呼吸引导、抽卡动画、情绪卡结果', ['呼吸引导：以短时、低压力的节奏帮助用户进入状态','抽取情绪卡：随机获得探索欲、创造欲、疲惫感、孤独感等主题','情绪解读：用温和语言解释当下感受，不做绝对判断','灵感寄语：提供一句可执行、可收藏的行动提示'], '进入情绪探索 → 完成呼吸引导 → 抽取情绪卡 → 阅读解读与寄语 → 收藏或返回探索')
add_section(doc,6,'灵感热度速览','通过简洁的数据反馈，让用户快速了解自己的灵感热度与探索轨迹，在看见变化的同时获得持续使用的成就感。','图片占位｜热度数值、趋势图、成就反馈', ['热度总览：展示近期灵感互动与收藏表现','趋势变化：用简洁图表呈现上升、平稳或回落','成就反馈：以里程碑或徽章记录持续探索','数据表达：避免复杂指标，让用户一眼读懂并愿意继续探索'], '进入个人空间 → 查看灵感热度 → 阅读趋势与成就 → 返回探索或继续收藏')

doc.add_page_break(); doc.add_heading('图片插入建议', level=1)
add_bullets(doc,['每个功能章节建议插入 1–3 张图：主界面截图 + 关键状态截图。','图片下方可补充一句图注，说明“用户在这里做什么”。','优先使用同一设备比例与同一主题，保证整份文档的视觉一致性。','若使用流程图，可按照“进入 → 操作 → 反馈 → 下一步”的顺序排列。'])
doc.add_heading('交付说明', level=2); doc.add_paragraph('本文档已预留图片占位框。你可以在 Word 中点击占位框并插入截图，也可以直接替换占位框所在表格。建议最终导出 PDF 前检查图片清晰度、页码和主题色一致性。')

doc.save(OUT)
print(OUT)
